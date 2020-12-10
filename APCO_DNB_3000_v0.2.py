#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jul 27 13:15:49 2020

@author: joshuagruen
"""

import tkinter as tk
import requests
from bs4 import BeautifulSoup
import pandas as pd
from datetime import date
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
import tkinter as tk
import os
import subprocess, os, platform
from PIL import ImageTk, Image
notaURLpopupused = False

def add_hyperlink(paragraph, text, url):
# This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.underline = True
    r.font.bold = True
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.name = 'Calibri'

    return hyperlink

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None





def notaURLpopup():
    notaURLpopup = tk.Tk()
    notaURLpopup.geometry("220x110+150+200")
    notaURLlabel = tk.Label(notaURLpopup, text = "Uh oh! We didn't like that URL. \nPlease try again.")
    notaURLlabel.pack()
    dismissnotaURL = tk.Button(notaURLpopup, text = "Dismiss", command = notaURLpopup.destroy)
    dismissnotaURL.pack()

   
def noselectionpopup():
   noselectionpopup = tk.Tk()
   noselectionpopup.geometry("275x100+150+200")
   noselectionlabel = tk.Label(noselectionpopup, text = "Uh oh! \nIt seems you forgot to select a category. \nPlease try again.")
   noselectionlabel.pack()
   dismissnoselection = tk.Button(noselectionpopup, text = "Dismiss", command = noselectionpopup.destroy)
   #command = dismissnoselection.destroy
   dismissnoselection.pack()
   
   
def urlpull(*args):
   
   paper = ""
   datesearchparent = ""
   datesearchchild = ""
   itlesearchparent = ""
   titlesearchchild = ""
   thedate = ""
   articlecat = ""
   manualyn = "n"
   continueyn = "1"
   thisday =""
   recognizedpaper = False
   papernickname = "null"
   global notaURLpopupused

   
   print("Feed me URL please")
   global DataFramePlaceholder
   URL = url_entry.get()
   print(URL)
   headers = { "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:73.0) Gecko/20100101 Firefox/73.0" }
   try:
       r1 = requests.get(URL,headers = headers)
       page  = r1.content
       soup1 = BeautifulSoup(page, 'html5lib')
   
   except:
       notaURLpopup()
       notaURLpopupused = True
       
   
#assign paper and search terms   
   
        #New York times
   if "nytimes.com" in URL:
    #if "www.nytimes.com" in URL:
            papernickname = "NYT"
            datesearchparent = 'meta'
            datesearchchild = {"property":"article:published"}
            titlesearchparent = "meta"
            titlesearchchild = {"property":"og:title"}
            recognizedpaper = True
            
       #Washington Post     
   elif "washingtonpost.com" in URL: 
            papernickname = "WaPo"
            titlesearchparent = 'meta'
            titlesearchchild = {"property": "og:title"}
            datesearchparent = 'meta'
            datesearchchild = {'name':'last_updated_date'}
            recognizedpaper = True
            
       
       #Times of Israel
   elif "timesofisrael.com" in URL:
            papernickname = "TOI"
            datesearchparent = "meta"
            datesearchchild = {'name':''}
            titlesearchparent = "meta"
            titlesearchchild = {'property':"og:title"}
            recognizedpaper = True
            
   elif "wsj.com" in URL:
            papernickname = "WSJ"
            titlesearchparent = "meta"
            titlesearchchild = {'property':'og:title'}
            recognizedpaper = True
            #bigtitle = soup1.title.string
            #print(bigtitle)
            #title, unused = bigtitle.split(' - ')
            #print(title+"part7b")
   
   elif "haaretz.com" in URL:
           titlesearchparent = "meta"
           titlesearchchild = {'property':'og:title'}
           papernickname = "Haaretz"
           recognizedpaper = True
           
   elif "jpost.com" in URL:
           titlesearchparent = "meta"
           titlesearchchild = {'property':'og:title'}
           papernickname = "JPost"
           recognizedpaper = True      
           
   elif "globes.co.il" in URL:
           titlesearchparent = "meta"
           titlesearchchild = {'property':'og:title'}
           papernickname = "Globes"
           recognizedpaper = True    
           
   elif "calcalistech.com" in URL:
           titlesearchparent = "meta"
           titlesearchchild = {'property':'og:title'}
           papernickname = "CTech"
           recognizedpaper = True   
           
      
  #try to find title of unknown website     
   else:
            try:
                #messytitle = soup1.find('h1')
                #title = messytitle.get_text()
                #print(title)
                
          #best option to find the title listed      
                title = soup1.find("meta", {"property":"og:title"})
                title = title["content"]
                print(title+"first part")
                
                if ' | ' in title:
                    title, unused = title.split(' | ')
                    
                elif ' - ' in title:
                    title, unused = title.split(' - ')
                    
                else:
                   title = title
                
                #sitename = soup1.find("meta", {'property':'og:site_name'})
                #print(sitename)
                #sitename = sitename["content"]
                #print(sitename)
                
                #author = soup1.find("meta", {'property':'og:author'})
                #print(author)
                
            except Exception as e:
                print(e)
                try:
                    bigtitle = soup1.title.string
                    print(title+"second part")
                    title, unused = bigtitle.split(' | ')
                    print(title+"part7a")
                except:
                        try:
                            bigtitle = soup1.title.string
                            title, unused = bigtitle.split(' - ')
                            print(title+"part7b")
                        except:
                            if notaURLpopupused == False:
                                notaURLpopup
                            else:
                                pass
            
            #if papernickname == "WSJ" or papernickname == "NYT":
               #pass
            #else:
               #papernickname = title
               
                    
            #try:
                #pull site name/paper
                #paper = soup1.find("meta", {'property':'og:site_name'})
                #print(paper)
                #paper = paper["content"]
                #print(paper+"third part")
                
            #except:
            try:
               bigsitename = soup1.title.string
               titleunused, paper = bigsitename.split(' | ')
               print(bigsitename+"fourthpart")
               print(paper+"fourth part")
            except:
                   try:
                       bigsitename = soup1.title.string
                       titleunused, paper = bigsitename.split(' - ')
                       print(paper+"fifth part")
                   except:
                       try:
                           paper = soup1.find('meta', {'property':'og:site_name'})
                           paper = paper["content"]
                           print(paper+"part9b")
                       except:
                           print("We are sorry. The paper you are looking for could not be found. \n Please enter it here.")
                           paper = input()
                           recognizedpaper = False
            
            
            print(paper+"sixth part")
            papernickname = paper
        
      
        
        #skip if manual entry
        #pull title
   if recognizedpaper == True:
        #pull title
           
               
               
                
            #pull date if has one requirement
            
            #if datesearchchild ==  "null":
                #messydate = soup1.find(datesearchparent)
                #thedate = messydate.get_text()
                #print(messydate)
                #print(thedate)
                #print("Date error!!")
                
           #pull date if has two requirements
            
            #else:
                #messydate = soup1.find(datesearchparent, datesearchchild)
                #print(messydate)
                #thedate = messydate.get_text()
                
            if titlesearchchild ==  "null":
                messytitle = soup1.find(titlesearchparent)
                title = messytitle.get_text('content')
                
                #print("Date error!!")
            
 
        
            else:
                messytitle = soup1.find(titlesearchparent, titlesearchchild)
                print(messytitle)
                title = messytitle['content']
                #print(soup1)
                
                
                
       #fix if says "today"
            if "Today" in thedate or "Today," in thedate or "today" in thedate or "today," in thedate:
                today = date.today()
                thisday =  str(today.strftime("%B %d, %Y"))
                thedate = thisday
                print(thedate)
                
            else:
               print(thedate)
                
           # if "Today," in date:
            #    date = date.today()
             #   print(date)
                
            #else:
           #    print(date)
                
                #print("Date error!!")
                
            #pullname
           # try: 
                #messyauthor = soup1.find_all('a', {'class':'byline'})
               # print(messyauthor)
                #author = messyauthor.get_text()
                #print(author)
            
           # except:
                #print("author error")
                
        
                                    
   else:
            print("Thank you!")
            
   print("\nIs this story: \n1. Domestic Affairs \n2. Internation Affairs \n3. Health \n4. Financial Services \n5. Tech \n6. Energy")
        
   
   global domcounter
   global intcounter
   global heacounter
   global fincounter
   global teccounter
   global enecounter

        
   articlecategorynumber = articlecategorynumbertk.get()
   print(str(articlecategorynumber) + "is the article category number")     
   if articlecategorynumbertk.get() == 1:
            articlecat = "Domestic Affairs"
            domcounter = domcounter+1
            
            
   elif articlecategorynumbertk.get() == 2:
            articlecat = "International Affairs"
            intcounter = intcounter+1
            
   elif articlecategorynumbertk.get() == 3:
            articlecat = "Health"
            heacounter = heacounter+1
            
   elif articlecategorynumbertk.get() == 4:
            articlecat = "Financial Services"
            fincounter = fincounter+1
            
   elif articlecategorynumbertk.get() == 5:
            articlecat = "Tech"
            teccounter = teccounter+1
            
   elif articlecategorynumbertk.get() == 6:
            articlecat = "Energy"
            enecounter = enecounter+1
            
   else:
       noselectionpopup()
            
   print("Category chosen: "+articlecat)
   print("Domestic Affairs: "+str(domcounter))
   print("International Affairs: "+str(intcounter))
   print("Health: "+str(heacounter))
   print("Financial Services: "+str(fincounter))
   print("Tech: "+str(teccounter))
   print("Energy: "+str(enecounter))
   domcountertk.set("Domestic Affairs: "+str(domcounter))
   intcountertk.set("International Affairs: "+str(intcounter))
   heacountertk.set("Health: "+str(heacounter))
   fincountertk.set("Financial Services: "+str(fincounter))
   teccountertk.set("Tech: "+str(teccounter))
   enecountertk.set("Energy: "+str(enecounter))
   
        
        #save data
   mostrecentarticletk.set("Added article: "+title[0:40]+"...")
   DataFramePlaceholder = DataFramePlaceholder.append({'stor_title': title, 'stor_paper': paper, 'stor_thedate': thedate, 'stor_articlecat': articlecat, 'stor_articlecategorynumber': articlecategorynumber, 'stor_URL': URL, 'stor_papernickname': papernickname}, ignore_index=True)
   pd.set_option("display.max_columns", None)
   print(DataFramePlaceholder)   
        #Ask if wants to continue
   print("\nWould you like to do another article? \n1. Yes \n2. No")
   continueyn = "1"
   url_entry.delete(0,'end')
    
def popupmsg():
    popup = tk.Tk()
    label = tk.Label(popup, text = "Are you sure you want to print?")
    label.pack()
    yesprintbutton = tk.Button(popup, text = "Yes", command = dataprinter, fg='red')
    yesprintbutton.pack()
    noprintbutton = tk.Button(popup, text = 'No', command = popup.destroy)
    noprintbutton.pack()
    
    
    
    
def dataprinter():
    global DataFramePlaceholder
    print("Your data just printed")
    #print(DataFramePlaceholder)
    #filename = input("Please choose a file name fam.\n")
    #filename = filename +  ".csv"
    #print(filename)
    
    today = date.today()
    thisday =  str(today.strftime("%B %d, %Y"))
    thedate = thisday
    print(thedate)
    
 #open template   
    dailynewsdoc = docx.Document("./Template for Python News Digest.docx")
    #print("Domestic Affairs")
    
    for paragraph in dailynewsdoc.paragraphs:
        if "***Date Goes Here***" in paragraph.text:
            datetoprint = paragraph.insert_paragraph_before("APCO Tel Aviv News Digest: "+thedate, style = 'APCO News Digest Header')
            delete_paragraph(paragraph)
            
            
 #Go through each section
    categorieslist = ['Domestic Affairs', 'International Affairs', 'Health', 'Financial Services', 'Tech', 'Energy']
    for x in categorieslist:
        #headerpara = dailynewsdoc.add_paragraph()
        from docx.shared import Pt
        entrycounter = 0
        #headerrun = headerpara.add_run(x)
        #headerrun.bold= True
        #headerrun.font.highlight_color = WD_COLOR_INDEX.YELLOW
        #headerrun.font.name =  'Arial'
        dacounter = 0
        
        #for i in range(len(DataFramePlaceholder)):
            #if DataFramePlaceholder.loc[i,'stor_articlecat'] == x:
                #print(DataFramePlaceholder.loc[i,'stor_title'])
                #dapara = dailynewsdoc.add_paragraph(DataFramePlaceholder.loc[i,'stor_title'], style ='List Bullet')
                #dapara = dailynewsdoc.add_paragraph("", style= 'List Bullet')
                #add_hyperlink(dapara, DataFramePlaceholder.loc[i,'stor_title'], DataFramePlaceholder.loc[i,'stor_URL'])
                #dapara.add_run(" ("+DataFramePlaceholder.loc[i,'stor_paper']+")").italic  = True
                
                #dacounter = dacounter+1
        
                

      #use key to find where to place text  
        findkey = ("***"+x+" Goes Here***")
        #go through each stored entry to see if applies to current section
        for i in range(len(DataFramePlaceholder)):
            if DataFramePlaceholder.loc[i,'stor_articlecat'] == x:
                entrycounter = entrycounter+1
                
                
                for paragraph in dailynewsdoc.paragraphs:
                    if findkey in paragraph.text:
                        bullet = paragraph.insert_paragraph_before("", style ='APCO News Digest')
                        #dapara = dailynewsdoc.add_paragraph(DataFramePlaceholder.loc[i,'stor_title'], style ='List Bullet')
                        #dapara = dailynewsdoc.add_paragraph("", style = 'List Bullet')
                        add_hyperlink(bullet, DataFramePlaceholder.loc[i,'stor_title'], DataFramePlaceholder.loc[i,'stor_URL'])
                        bullet.add_run(" ("+DataFramePlaceholder.loc[i,'stor_papernickname']+")").italic  = True
                        #delete_paragraph(paragraph)
                    
                    
        if entrycounter == 0:
            for paragraph in dailynewsdoc.paragraphs:
                if findkey in paragraph.text:
                    noresults = paragraph.insert_paragraph_before("", style ='APCO News Digest')
                        #dapara = dailynewsdoc.add_paragraph(DataFramePlaceholder.loc[i,'stor_title'], style ='List Bullet')
                        #dapara = dailynewsdoc.add_paragraph("", style = 'List Bullet')
                    noresults.add_run("There is no news to report.").italic  = True
                    #delete_paragraph(paragraph)
                        
    for x in categorieslist:
        findkey = ("***"+x+" Goes Here***")
        for paragraph in dailynewsdoc.paragraphs:
            if findkey in paragraph.text:
                delete_paragraph(paragraph)
            
    
    docnameforsave = ("../Daily News Briefing "+str(thedate)+".docx")
    dailynewsdoc.save(docnameforsave)
    
    if platform.system() == 'Darwin':       # macOS
        subprocess.call(('open', docnameforsave))
                        
    elif platform.system() == 'Windows':
        os.startfile((docnameforsave))
        
    else:                                   # linux variants
        subprocess.call(('xdg-open', docnameforsave))
    
    #DataFramePlaceholder.to_csv(filename)
        
        
window = tk.Tk()
window.title("APCO DNB 3000")
window.geometry("398x290+75+75")
window.resizable(False, False)
greeting = tk.Label(window, text="Insert URL here: ", font='bold')
#domcounter = ("Number of Cases: "+str(five))
var = tk.StringVar(window)
domcountertk = tk.StringVar(window)
intcountertk = tk.StringVar(window)
heacountertk = tk.StringVar(window)
fincountertk = tk.StringVar(window)
teccountertk = tk.StringVar(window)
enecountertk = tk.StringVar(window)
mostrecentarticletk=tk.StringVar(window)
domcountertk.set("Domestic Affairs: 0")
intcountertk.set("International Affairs: 0")
heacountertk.set("Health: 0")
fincountertk.set("Financial Services: 0")
teccountertk.set("Tech: 0")
enecountertk.set("Energy: 0")
logoimage = ImageTk.PhotoImage(Image.open("./APCO DND 3000 Logo 3.jpg"),master=window)
logotk = tk.Label(window, image = logoimage)

counterbgcolor = 'DodgerBlue3'
countersframe = tk.Frame(window, bg = counterbgcolor)

#setcounters and color
countersframe.grid(row=6, column=0, columnspan=4,rowspan=4, sticky='ew')
domcounter_label = tk.Label(countersframe, textvariable = domcountertk, bg=counterbgcolor, fg='white')
intcounter_label = tk.Label(countersframe, textvariable = intcountertk, bg=counterbgcolor, fg='white')
heacounter_label = tk.Label(countersframe, textvariable = heacountertk, bg=counterbgcolor, fg='white')
fincounter_label = tk.Label(countersframe, textvariable = fincountertk, bg=counterbgcolor, fg='white')
teccounter_label = tk.Label(countersframe, textvariable = teccountertk, bg=counterbgcolor, fg='white')
enecounter_label = tk.Label(countersframe, textvariable = enecountertk, bg=counterbgcolor, fg='white')
mostrecentarticle_label=tk.Label(countersframe, textvariable=mostrecentarticletk, bg=counterbgcolor, fg='white')
#create buttons
articlecategorynumbertk = tk.IntVar(window)
domestic_button = tk.Radiobutton(window, text="Domestic Affairs", variable = articlecategorynumbertk, value = 1)
international_button = tk.Radiobutton(window, text="International Affairs", variable = articlecategorynumbertk, value = 2)
health_button = tk.Radiobutton(window, text="Health", variable = articlecategorynumbertk, value = 3)
financial_button = tk.Radiobutton(window, text="Financial Services", variable = articlecategorynumbertk, value = 4)
tech_button = tk.Radiobutton(window, text="Tech", variable = articlecategorynumbertk, value = 5)
energy_button = tk.Radiobutton(window, text="Energy", variable = articlecategorynumbertk, value = 6)


domcounter = 0
intcounter = 0
heacounter = 0
fincounter = 0
teccounter = 0
enecounter = 0
DataFramePlaceholder = pd.DataFrame(columns=['stor_title','stor_paper','stor_thedate','stor_articlecat','stor_articlecategorynumber','stor_URL', 'stor_papernickname'])


#create URL entry and functional buttons
url_entry = tk.Entry(window, text = "Enter URL here", width = 35)
add_article_button = tk.Button(window, text = "Add URL", bg = "blue", command = urlpull)#)
print_button = tk.Button(window, text = "Print", command = popupmsg, height=2, width=15)
#url_entry.bind(add_article_button, urlpull)
url_entry.bind("<Return>", urlpull)




var_label_check = tk.Label(window, fg="black", bg = "green", width = 50, textvariable=var)





logotk.grid(row=1, column=0, columnspan=3)

greeting.grid(row=2, column=0, columnspan=3)
url_entry.grid(row=3, column=0,columnspan=2)
add_article_button.grid(row=3, column=2)

#var_label_check.pack()
firstbuttonrow  = 4
domestic_button.grid(row=firstbuttonrow, column=0,sticky='W')
international_button.grid(row=firstbuttonrow, column=1,sticky='W')
health_button.grid(row=firstbuttonrow, column=2, sticky='W')
secondbuttonrow = 5
financial_button.grid(row=secondbuttonrow, column=0, sticky='W')
tech_button.grid(row=secondbuttonrow,column=1,sticky='W')
energy_button.grid(row=secondbuttonrow, column=2,sticky='W')

mostrecentarticle_label.grid(row=6,column=0, columnspan=3, sticky='W')
firstcounterrow=7
domcounter_label.grid(row=firstcounterrow, column=0, sticky='W')
intcounter_label.grid(row=firstcounterrow, column=1, sticky='W')
heacounter_label.grid(row=firstcounterrow, column=2, sticky='W')
secondcounterrow=8
fincounter_label.grid(row=secondcounterrow, column=0, sticky='W')
teccounter_label.grid(row=secondcounterrow, column=1, sticky = 'W')
enecounter_label.grid(row=secondcounterrow, column=2, sticky='W')


print_button.place(bordermode="outside", x=130, y=242)
